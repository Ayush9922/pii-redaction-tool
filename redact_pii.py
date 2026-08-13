import re
import json
import os
from faker import Faker
from docx import Document
import docx.shared

class PIIRedactor:
    def __init__(self):
        self.fake = Faker()
        Faker.seed(42)
        # Maintain consistent mapping across the run
        self.mappings = {}

    def get_fake(self, original, pii_type):
        key = (original.strip(), pii_type)
        if key in self.mappings:
            return self.mappings[key]

        if pii_type == 'email':
            fake_val = self.fake.email()
        elif pii_type == 'phone':
            digits = "".join(str(self.fake.random_digit()) for _ in range(10))
            fake_val = f"+91 {digits}"
        elif pii_type == 'ssn':
            fake_val = self.fake.ssn()
        elif pii_type == 'credit_card':
            # Return a 16-digit hyphen-separated number
            raw = self.fake.credit_card_number()
            # Format as XXXX-XXXX-XXXX-XXXX
            raw = raw.replace(' ', '').replace('-', '')[:16]
            fake_val = f"{raw[:4]}-{raw[4:8]}-{raw[8:12]}-{raw[12:16]}"
        elif pii_type == 'ip':
            fake_val = self.fake.ipv4()
        elif pii_type == 'dob':
            fake_date = self.fake.date_of_birth(minimum_age=25, maximum_age=65)
            parts = original.split('-')
            if len(parts[0]) == 4:
                fake_val = fake_date.strftime('%Y-%m-%d')
            else:
                fake_val = fake_date.strftime('%d-%m-%Y')
        elif pii_type == 'company':
            fake_val = self.fake.company()
        elif pii_type == 'address':
            fake_val = self.fake.address().replace('\n', ', ')
        elif pii_type == 'name':
            fake_val = self.fake.name()
        else:
            fake_val = f"[REDACTED_{pii_type.upper()}]"

        self.mappings[key] = fake_val
        return fake_val


# ---------------------------------------------------------------------------
# PII detection helpers
# ---------------------------------------------------------------------------

# Words that disqualify a two-word capitalised phrase from being a person name
_NON_NAME_WORDS = {
    "ticket", "date", "client", "company", "subject", "support", "team", "ops",
    "finance", "systems", "consulting", "biotech", "solutions", "group", "ltd",
    "inc", "corp", "corporation", "road", "street", "terrace", "lane", "avenue",
    "drive", "boulevard", "sector", "mg", "camp", "park", "west", "north", "south",
    "east", "evergreen", "springfield", "india", "usa", "hello", "thanks", "regards",
    "dear", "hi", "best", "reset", "billing", "alternate", "password",
    "api", "connection", "refused", "payment", "issue", "mastercard", "visa",
    "refusal", "agent", "customer", "verification", "account", "please", "update",
    "could", "check", "will", "contact", "need", "tried", "system", "charge",
    "transaction", "failed", "again", "escalated", "network", "access", "review",
    "database", "error", "connection", "refused", "production", "environment"
}

def _is_likely_name(text: str) -> bool:
    words = text.lower().split()
    for w in words:
        if w in _NON_NAME_WORDS:
            return False
    return True


def detect_pii(text: str):
    """
    Returns a list of (start, end, pii_type, matched_string) tuples,
    non-overlapping, sorted by start index.
    """
    raw_spans = []

    # --- Priority 1: SSN (most specific format) ---
    for m in re.finditer(r'\b\d{3}-\d{2}-\d{4}\b', text):
        raw_spans.append((m.start(), m.end(), 'ssn', m.group()))

    # --- Priority 2: Credit card (16 digits, optional separators) ---
    for m in re.finditer(r'\b\d{4}[- ]?\d{4}[- ]?\d{4}[- ]?\d{4}\b', text):
        raw_spans.append((m.start(), m.end(), 'credit_card', m.group()))

    # --- Priority 3: IP address ---
    for m in re.finditer(r'\b(?:\d{1,3}\.){3}\d{1,3}\b', text):
        raw_spans.append((m.start(), m.end(), 'ip', m.group()))

    # --- Priority 4: Email ---
    for m in re.finditer(r'\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b', text):
        raw_spans.append((m.start(), m.end(), 'email', m.group()))

    # --- Priority 5: Phone ---
    for m in re.finditer(r'\+\d{1,3}\s?\d{5,10}\b', text):
        raw_spans.append((m.start(), m.end(), 'phone', m.group()))

    # --- Priority 6: DOB  ---
    # YYYY-MM-DD (years 1900-2009 only, to avoid matching 2026-08-10 ticket dates)
    for m in re.finditer(r'\b(?:19\d{2}|200\d)-(?:0[1-9]|1[0-2])-(?:0[1-9]|[12]\d|3[01])\b', text):
        raw_spans.append((m.start(), m.end(), 'dob', m.group()))
    # DD-MM-YYYY
    for m in re.finditer(r'\b(?:0[1-9]|[12]\d|3[01])-(?:0[1-9]|1[0-2])-(?:19\d{2}|200\d)\b', text):
        raw_spans.append((m.start(), m.end(), 'dob', m.group()))

    # --- Priority 7: Address (India style) ---
    for m in re.finditer(
        r'\d+,\s*[A-Za-z0-9 .]+,\s*[A-Za-z0-9 .]+,\s*[A-Za-z0-9 .]+,\s*[A-Za-z0-9 .]+,\s*\d{5,6},\s*India',
        text, re.IGNORECASE
    ):
        raw_spans.append((m.start(), m.end(), 'address', m.group()))

    # US address style: digits + street + city + state abbr + zip + country
    for m in re.finditer(
        r'\b\d+\s+[A-Za-z ]+(?:Terrace|Street|Road|Avenue|Lane|Drive|Boulevard|Way|Court|Plaza|St|Rd|Ave|Ln|Dr|Ct|Plz|Blvd),\s*[A-Za-z ]+,\s*[A-Z]{2}\s*\d{5},\s*USA\b',
        text, re.IGNORECASE
    ):
        raw_spans.append((m.start(), m.end(), 'address', m.group()))

    # --- Priority 8: Company ---
    for m in re.finditer(
        r'\b[A-Z][A-Za-z0-9]+(?: [A-Z][A-Za-z0-9]+)* (?:Ltd\.|Inc\.|Co\.|Systems|Solutions|Group|Consulting|Biotech|Corporation|Corp)\b',
        text
    ):
        raw_spans.append((m.start(), m.end(), 'company', m.group()))

    # --- Priority 9: Person name (Two Title-Case words) ---
    for m in re.finditer(r'\b([A-Z][a-z]+) ([A-Z][a-z]+)\b', text):
        val = m.group()
        if _is_likely_name(val):
            raw_spans.append((m.start(), m.end(), 'name', val))

    # --- Resolve overlaps: prefer longer/higher-priority spans ---
    # Sort: start asc, end desc (so longer spans win ties at same start)
    raw_spans.sort(key=lambda x: (x[0], -x[1]))

    resolved = []
    last_end = -1
    for start, end, pii_type, value in raw_spans:
        if start >= last_end:
            resolved.append((start, end, pii_type, value))
            last_end = end

    return resolved


def redact_text(text: str, redactor: PIIRedactor):
    """
    Detect PII in text and replace with consistent fake values.
    Returns (redacted_text, list_of_replacement_dicts).
    """
    spans = detect_pii(text)
    # Replace from right to left to preserve indices
    spans_rev = sorted(spans, key=lambda x: x[0], reverse=True)

    redacted = text
    replacements = []
    for start, end, pii_type, value in spans_rev:
        fake_val = redactor.get_fake(value, pii_type)
        redacted = redacted[:start] + fake_val + redacted[end:]
        replacements.append({
            "start": start,
            "end": end,
            "type": pii_type,
            "original": value,
            "replacement": fake_val
        })

    return redacted, replacements


def build_docx(redacted_content: str, output_path: str):
    doc = Document()

    # Title
    title = doc.add_heading("Redacted Ticket Log", level=0)

    header_prefixes = ("Ticket ID:", "Date:", "Client:", "Company:", "Subject:")
    separator_prefixes = ("===", "---")

    for line in redacted_content.split('\n'):
        stripped = line.strip()
        if any(stripped.startswith(p) for p in header_prefixes):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
        elif any(stripped.startswith(p) for p in separator_prefixes):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = docx.shared.Pt(6)
            p.paragraph_format.space_after = docx.shared.Pt(6)
            p.add_run(stripped)
        else:
            doc.add_paragraph(stripped)

    doc.save(output_path)


def main():
    input_path = "ticket_log.txt"
    output_docx = "redacted_ticket_log.docx"
    output_txt  = "redacted_ticket_log.txt"
    audit_path  = "redaction_mappings.json"

    if not os.path.exists(input_path):
        print(f"[!] Error: '{input_path}' not found.")
        return

    with open(input_path, 'r', encoding='utf-8') as f:
        original = f.read()

    redactor = PIIRedactor()
    redacted, replacements = redact_text(original, redactor)

    # ---- Save plain-text redacted version ----
    with open(output_txt, 'w', encoding='utf-8') as f:
        f.write(redacted)
    print(f"[+] Saved redacted text  -> {output_txt}")

    # ---- Save Word document ----
    build_docx(redacted, output_docx)
    print(f"[+] Saved redacted .docx -> {output_docx}")

    # ---- Save audit log ----
    # Convert tuple keys to strings for JSON serialisation
    json_mappings = {f"{k[0]} [{k[1]}]": v for k, v in redactor.mappings.items()}
    with open(audit_path, 'w', encoding='utf-8') as f:
        json.dump({"mappings": json_mappings, "detections": replacements}, f, indent=2)
    print(f"[+] Saved audit log      -> {audit_path}")

    print(f"\n[Summary] {len(replacements)} PII instances redacted across {len(redactor.mappings)} unique entities.")


if __name__ == '__main__':
    main()
